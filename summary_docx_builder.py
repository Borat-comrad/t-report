from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

from data_quality import SourceRowIssue
from email_body_builder import build_key_findings
from entity_mapper import DeliveryItem
from kp_analyzer import KpReportItem, build_request_state
from management_insights import BottleneckSummary, ManagerSummary
from report_metric_catalog import (
    DOCX_BOTTLENECK_HEADERS,
    DOCX_DELIVERY_HEADERS,
    DOCX_KP_HEADERS,
    DOCX_MANAGER_HEADERS,
    SUMMARY_ROW_LABELS,
)
from weekly_report_models import BuiltWeeklyReport


BLACK = RGBColor(0, 0, 0)


def format_date_value(value: date | None) -> str:
    if value is None:
        return ""

    return value.strftime("%d.%m.%Y")


def format_float_value(value: float | None) -> str:
    if value is None:
        return ""

    return f"{value:.2f}"


def format_avg_value(value: float | None) -> str:
    if value is None:
        return ""

    return f"{value:.1f}"


def apply_document_styles(document: Document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)
    normal_style.font.color.rgb = BLACK

    for style_name, font_size in (("Heading 1", 16), ("Heading 2", 13)):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(font_size)
        style.font.bold = True
        style.font.color.rgb = BLACK


def append_summary_table(document: Document, report: BuiltWeeklyReport) -> None:
    summary = report.summary

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    header_cells[0].text = "Показатель"
    header_cells[1].text = "Значение"

    rows = [
        (SUMMARY_ROW_LABELS[0], format_date_value(summary.report_date)),
        (SUMMARY_ROW_LABELS[1], str(summary.critical_kp_count)),
        (SUMMARY_ROW_LABELS[2], str(summary.warning_kp_count)),
        (SUMMARY_ROW_LABELS[3], str(summary.current_week_delivery_count)),
        (SUMMARY_ROW_LABELS[4], str(summary.next_two_weeks_delivery_count)),
        (SUMMARY_ROW_LABELS[5], str(len(report.manager_summaries))),
    ]

    for label, value in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value


def append_key_findings(document: Document, report: BuiltWeeklyReport) -> None:
    document.add_heading("Ключевые наблюдения", level=2)

    for finding in build_key_findings(report):
        document.add_paragraph(finding, style="List Bullet")


def append_manager_summary_table(document: Document, items: list[ManagerSummary]) -> None:
    document.add_heading("Сводка по каждому менеджеру", level=2)

    if not items:
        document.add_paragraph("Менеджеры в зоне контроля не выявлены.")
        return

    table = document.add_table(rows=1, cols=9)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for index, value in enumerate(DOCX_MANAGER_HEADERS):
        header_cells[index].text = value

    for item in items:
        row_cells = table.add_row().cells
        row_cells[0].text = item.display_name
        row_cells[1].text = str(item.critical_kp_count)
        row_cells[2].text = str(item.warning_kp_count)
        row_cells[3].text = format_avg_value(item.avg_overdue_days)
        row_cells[4].text = str(item.current_week_delivery_count)
        row_cells[5].text = str(item.next_two_weeks_delivery_count)
        row_cells[6].text = str(item.workload_score)
        row_cells[7].text = item.efficiency_assessment
        row_cells[8].text = item.bottleneck_reason


def append_bottleneck_table(
    document: Document,
    title: str,
    items: list[BottleneckSummary],
    limit: int = 10,
) -> None:
    document.add_heading(title, level=2)

    if not items:
        document.add_paragraph("Записи отсутствуют.")
        return

    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for index, value in enumerate(DOCX_BOTTLENECK_HEADERS):
        header_cells[index].text = value

    for item in items[:limit]:
        row_cells = table.add_row().cells
        row_cells[0].text = item.group_name
        row_cells[1].text = str(item.overdue_kp_count)
        row_cells[2].text = str(item.critical_kp_count)
        row_cells[3].text = str(item.warning_kp_count)
        row_cells[4].text = format_avg_value(item.avg_overdue_days)
        row_cells[5].text = item.affected_managers
        row_cells[6].text = item.diagnostic_comment


def append_kp_top_table(document: Document, items: list[KpReportItem], limit: int = 5) -> None:
    document.add_heading("Топ критичных КП", level=2)

    if not items:
        document.add_paragraph("Критичные КП отсутствуют.")
        return

    table = document.add_table(rows=1, cols=8)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for index, value in enumerate(DOCX_KP_HEADERS):
        header_cells[index].text = value

    for item in items[:limit]:
        row_cells = table.add_row().cells
        row_cells[0].text = item.request_number
        row_cells[1].text = format_date_value(item.received_date)
        row_cells[2].text = item.manufacturer
        row_cells[3].text = item.branch
        row_cells[4].text = item.employee
        row_cells[5].text = format_date_value(item.proposal_deadline)
        row_cells[6].text = str(item.overdue_days)
        row_cells[7].text = item.request_state or build_request_state(item.overdue_days)


def append_delivery_table(
    document: Document,
    title: str,
    items: list[DeliveryItem],
    limit: int = 5,
) -> None:
    document.add_heading(title, level=2)

    if not items:
        document.add_paragraph("Записи отсутствуют.")
        return

    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for index, value in enumerate(DOCX_DELIVERY_HEADERS):
        header_cells[index].text = value

    for item in items[:limit]:
        row_cells = table.add_row().cells
        row_cells[0].text = item.code
        row_cells[1].text = item.ro_number
        row_cells[2].text = item.branch
        row_cells[3].text = format_float_value(item.amount_eur_without_vat)
        row_cells[4].text = format_date_value(item.delivery_date)
        row_cells[5].text = item.owner


def append_source_row_issues(document: Document, issues: list[SourceRowIssue]) -> None:
    document.add_heading("Строки, исключенные из расчета", level=2)

    if not issues:
        document.add_paragraph("Строк, исключенных из расчета, нет.")
        return

    document.add_paragraph(
        "Ниже перечислены строки исходной таблицы, которые не попали в расчет из-за ошибок в данных."
    )

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    header_cells[0].text = "Лист"
    header_cells[1].text = "Строка"
    header_cells[2].text = "Колонка"
    header_cells[3].text = "Значение"
    header_cells[4].text = "Что исправить"

    for issue in issues:
        row_cells = table.add_row().cells
        row_cells[0].text = issue.sheet_name
        row_cells[1].text = str(issue.row_index)
        row_cells[2].text = issue.column_name
        row_cells[3].text = issue.raw_value
        row_cells[4].text = issue.reason


def build_summary_docx_document(report: BuiltWeeklyReport) -> Document:
    document = Document()
    apply_document_styles(document)

    document.add_heading("Еженедельный отчет по КП и поставкам", level=1)
    document.add_paragraph(f"Дата отчета: {format_date_value(report.summary.report_date)}")

    append_summary_table(document, report)
    append_key_findings(document, report)
    append_manager_summary_table(document, report.manager_summaries)
    append_bottleneck_table(document, "Косвенные узкие места по производителям", report.manufacturer_bottlenecks)
    append_bottleneck_table(document, "Косвенные узкие места по филиалам", report.branch_bottlenecks)
    append_kp_top_table(document, report.critical_kp_items, limit=5)
    append_delivery_table(
        document,
        "Поставки на текущую неделю",
        report.current_week_delivery_items,
        limit=5,
    )
    append_delivery_table(
        document,
        "Поставки на следующие 2 недели",
        report.next_two_weeks_delivery_items,
        limit=5,
    )
    append_source_row_issues(document, report.source_row_issues)

    return document


def build_summary_docx_file(report: BuiltWeeklyReport, output_path: Path) -> Path:
    document = build_summary_docx_document(report)
    document.save(output_path)
    return output_path
