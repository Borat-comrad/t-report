from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

from report_metric_catalog import METRIC_REFERENCE_SECTIONS, MetricSection


BLACK = RGBColor(0, 0, 0)


def apply_document_styles(document: Document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = BLACK

    for style_name, font_size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(font_size)
        style.font.bold = True
        style.font.color.rgb = BLACK


def append_metric_section(document: Document, section: MetricSection) -> None:
    document.add_heading(section.title, level=2)
    document.add_paragraph(section.intro)

    for metric in section.metrics:
        document.add_heading(metric.name, level=3)
        table = document.add_table(rows=4, cols=2)
        table.style = "Table Grid"
        labels = (
            ("Что это", metric.description),
            ("Как считается", metric.formula),
            ("О чем говорит", metric.interpretation),
            ("Единица / формат", metric.unit),
        )
        for row_index, (label, value) in enumerate(labels):
            row = table.rows[row_index].cells
            row[0].text = label
            row[1].text = value

        document.add_paragraph("")


def build_metrics_reference_document() -> Document:
    document = Document()
    apply_document_styles(document)

    document.add_heading("Универсальная справка по метрикам отчета", level=1)
    document.add_paragraph(
        "Документ поясняет, что означает каждая ключевая метрика отчета, как она считается и какой управленческий вывод из нее можно сделать."
    )

    for section in METRIC_REFERENCE_SECTIONS:
        append_metric_section(document, section)

    return document


def build_metrics_reference_docx_file(output_path: Path) -> Path:
    document = build_metrics_reference_document()
    document.save(output_path)
    return output_path
